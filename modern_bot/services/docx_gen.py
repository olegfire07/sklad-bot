import asyncio
import random
import string
import logging
from typing import Dict, Any
from pathlib import Path
from datetime import datetime
from docx import Document
from docx.shared import Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from modern_bot.config import TEMPLATE_PATH, DOCS_DIR
from modern_bot.utils.files import sanitize_filename
from modern_bot.database.db import load_user_data

logger = logging.getLogger(__name__)

def replace_placeholders_in_document(doc: Document, placeholders: Dict[str, str]) -> None:
    def _replace_in_runs(runs):
        for run in runs:
            for key, value in placeholders.items():
                if key in run.text:
                    run.text = run.text.replace(key, value)

    for paragraph in doc.paragraphs:
        if paragraph.runs:
            _replace_in_runs(paragraph.runs)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if paragraph.runs:
                        _replace_in_runs(paragraph.runs)

def add_borders_to_table(table: Any) -> None:
    for row in table.rows:
        for cell in row.cells:
            tc = cell._element
            tcPr = tc.get_or_add_tcPr()
            borders = OxmlElement('w:tcBorders')
            for border in ['top', 'left', 'bottom', 'right']:
                border_element = OxmlElement(f"w:{border}")
                border_element.set(qn('w:val'), 'single')
                border_element.set(qn('w:sz'), '8')
                border_element.set(qn('w:space'), '0')
                border_element.set(qn('w:color'), 'auto')
                borders.append(border_element)
            tcPr.append(borders)

def populate_table_with_data(doc: Document, data: Dict[str, Any]) -> None:
    if not doc.tables:
        logger.error("No tables found in document.")
        return
    table = doc.tables[0]
    for i, item in enumerate(data.get('photo_desc', []), 1):
        try:
            new_row = table.add_row()
            row_cells = new_row.cells
            if len(row_cells) < 8:
                logger.error("Table structure mismatch (less than 8 columns).")
                continue

            photo_path = Path(item.get('photo', ""))
            row_cells[0].text = str(i)
            if photo_path.is_file():
                p = row_cells[2].paragraphs[0] if row_cells[2].paragraphs else row_cells[2].add_paragraph()
                p.add_run().add_picture(str(photo_path), width=Inches(1.0))
            else:
                row_cells[2].text = 'Фото отсутствует'

            description = item.get('description') or 'Нет описания'
            evaluation_value = item.get('evaluation') or 'Нет данных'
            row_cells[1].text = description
            row_cells[5].text = evaluation_value
            row_cells[6].text = evaluation_value
            row_cells[7].text = 'да'
        except Exception as e:
            logger.error(f"Error populating table: {e}")
    add_borders_to_table(table)

async def create_document(user_id: int, username: str = "") -> Path:
    data = await load_user_data(user_id)
    if not data:
        raise ValueError("No user data found.")
    if not TEMPLATE_PATH.exists():
        raise FileNotFoundError(f"Template '{TEMPLATE_PATH}' not found.")

    DOCS_DIR.mkdir(parents=True, exist_ok=True)
    selected_date = data.get('date') or datetime.now().strftime('%d.%m.%Y')
    timestamp = datetime.now().strftime('%H-%M-%S')
    placeholders = {
        '{date}': selected_date,
        '{issue_number}': data.get('issue_number', 'Не указано'),
        '{department_number}': data.get('department_number', 'Не указано'),
        '{region}': data.get('region', 'Не указано'),
        '{ticket_number}': data.get('ticket_number', 'Не указано'),
        '{username}': username
    }
    
    base_filename = (f"{placeholders['{department_number}']}, Заключение антиквариат № "
                     f"{placeholders['{issue_number}']} (билет {placeholders['{ticket_number}']}), "
                     f"{placeholders['{region}']}, от {selected_date} {timestamp}.docx")
    
    safe_filename_str = sanitize_filename(base_filename)
    if not safe_filename_str:
        safe_filename_str = f"Заключение_{timestamp}.docx"
    filepath = DOCS_DIR / safe_filename_str

    suffix = Path(safe_filename_str).suffix or ".docx"
    stem = Path(safe_filename_str).stem or "Заключение"
    while filepath.exists():
        unique_suffix = ''.join(random.choices(string.ascii_lowercase + string.digits, k=4))
        candidate_name = sanitize_filename(f"{stem}_{unique_suffix}{suffix}")
        if not candidate_name:
            candidate_name = f"Заключение_{timestamp}_{unique_suffix}.docx"
        filepath = DOCS_DIR / candidate_name
        safe_filename_str = candidate_name

    def _build_document():
        try:
            doc = Document(TEMPLATE_PATH)
            if doc.paragraphs:
                doc.paragraphs[0].insert_paragraph_before(filepath.stem)
            else:
                doc.add_paragraph(filepath.stem)
            replace_placeholders_in_document(doc, placeholders)
            populate_table_with_data(doc, data)
            doc.save(filepath)
        except Exception as doc_error:
            logger.error(f"Failed to build document {filepath}: {doc_error}", exc_info=True)
            raise

    try:
        await asyncio.to_thread(_build_document)
    except Exception as exc:
        raise RuntimeError("Error generating document.") from exc
    logger.info(f"Document saved: {filepath}")
    return filepath
